from docx import Document
from docx.shared import Inches
import os
from docx import Document


def create_word_report(user_id, buffer):
    # Papka mavjudligini tekshirish
    os.makedirs("temp_word", exist_ok=True)
    file_name = f"temp_word/report_{user_id}.docx"

    doc = Document()
    doc.add_heading('@Slayd_chi Telegram bot ', 0)

    for item in buffer:
        try:
            if item['type'] == 'text':
                # Bo'sh matn yuborilsa xato bermasligi uchun
                content = item.get('content', '').strip()
                if content:
                    doc.add_paragraph(content)

            elif item['type'] == 'image':
                img_path = item.get('content')
                if img_path and os.path.exists(img_path):
                    # Rasm o'lchamini Inches bilan berish
                    doc.add_picture(img_path, width=Inches(5.0))
                    # Rasm qo'shilgach, uni o'chirish (agar handle_collection'da o'chirilmasa)
                    # Lekin ehtiyot bo'ling: bu funksiya ichida o'chirilsa,
                    # cmd_save_doc funksiyasidagi finaly blokida xato berishi mumkin.
        except Exception as e:
            print(f"Xatolik yuz berdi: {e}")
            continue

    doc.save(file_name)
    return file_name
