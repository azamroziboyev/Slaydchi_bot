from docx import Document
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time


def edit_mustaqil(input_path: str, output_path: str, universitet: str, student_name: str, qabul_qiluvchi: str, year: str = None):
    """Edit the first page of an existing document with university, student and teacher text and save the result."""
    document = Document(input_path)
    
    if not year:
        year = time.strftime("%Y")  # current year

    # Update normal style
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(14)

    # Clear first page by removing all paragraphs (optional)
    while len(document.paragraphs) > 0:
        p = document.paragraphs[0]
        # remove paragraph contents to rebuild the cover page
        p.clear()

    # Rebuild first page
    p = document.add_paragraph("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM FAN VA INNOVATSIYALAR VAZIRLIGI")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph(universitet)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        document.add_paragraph("")

    p = document.add_paragraph()
    run = p.add_run('REFERAT')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        document.add_paragraph("")

    bajardi = document.add_paragraph()
    run = bajardi.add_run("Bajardi:")
    run.bold = True
    bajardi.add_run(f"                                                                    {student_name}")
    for _ in range(3):
        document.add_paragraph("")

    qabul_qildi = document.add_paragraph()
    run = qabul_qildi.add_run("Qabul qildi:")
    run.bold = True
    qabul_qildi.add_run(f"                                                              {qabul_qiluvchi}")
    document.add_paragraph("")
    document.add_paragraph("")

    yil = document.add_paragraph(year)
    yil.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # Save changes to output_path
    document.save(output_path)
    return output_path