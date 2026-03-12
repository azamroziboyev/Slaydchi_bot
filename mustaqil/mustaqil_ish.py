from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from dotenv import load_dotenv
import requests, json, time
from text import TRANSLATIONS

# ------------------------------
# 1. CREATE FIRST PAGE
# ------------------------------
def make_first_page(universitet, student_name, qabul_qiluvchi, lang):
    """Create and return the title page for an independent work (mustaqil ish).

    Uses a project-local template file and TRANSLATIONS for language-specific labels.
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    t = TRANSLATIONS[lang]
    # Go one level up to project root
    project_root = os.path.abspath(os.path.join(script_dir, ".."))

    # Build full path to the template
    template_path = os.path.join(project_root, "template_mustaqil.docx")
    document = Document(template_path)
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(14)
    p = document.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph(f"{universitet}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        document.add_paragraph("")

    p = document.add_paragraph()
    run = p.add_run(t["mustaqil_ishi_titul"])
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        document.add_paragraph("")

    bajardi = document.add_paragraph()
    run = bajardi.add_run(t["bajaruvchi_titul"])
    run.bold = True
    bajardi.add_run(f"                                                                    {student_name}")
    for _ in range(3):
        document.add_paragraph("")

    qabul_qildi = document.add_paragraph()
    run = qabul_qildi.add_run(t["qabul_qiluvchi_titul"])
    run.bold = True
    qabul_qildi.add_run(f"                                                              {qabul_qiluvchi}")
    document.add_paragraph("")
    document.add_paragraph("")

    year = time.strftime("%Y")
    yil = document.add_paragraph(f"{year}y")
    yil.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()
    return document


# ------------------------------
# 2. LOAD API KEYS
# ------------------------------
load_dotenv()
# key = os.getenv("BYTEZ_API_KEY")

DEEPSEEK_KEY = os.getenv("DEEPSEEK2")
Llama_key = os.getenv("LLAMA")
GROK_KEY = os.getenv("DEEPSEEK2")
PERPLEX_KEY = os.getenv("PERPLEX_KEY")
API_KEY = os.getenv("NOVA_KEY")

# ------------------------------
# 3. AI TEXT GENERATOR
# ------------------------------

#original ishlaydigani
def generate_deepseek(prompt):
    """Call Deepseek API to generate a short academic paragraph for prompt.

    Returns the generated string or "error" if the remote call fails.
    """
    if not DEEPSEEK_KEY:
        return "Bu mavzuga oid qisqacha akademik monolog."
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {DEEPSEEK_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "deepseek/deepseek-chat",
        "messages": [
            {"role": "system", "content": "You write clear academic texts."},
            {"role": "user", "content": f"{prompt}."}
        ]
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=12)
        response.raise_for_status()  # Raise an exception for bad status codes
        result = response.json()
        return result["choices"][0]["message"]["content"]
    except Exception as e:
        return "error"
# def generate_deepseek(prompt):
#     sdk = Bytez(key)
#     model = sdk.model("Qwen/Qwen3-4B-Instruct-2507")
#     response = model.run([
#         {
#             "role": "user",
#             "content": prompt
#         }
#     ])

#     if response.error:
#         raise RuntimeError(response.error)

#     return response.output["content"]


#original ishlaydigani
def expand_grok_ai(paragraph, target_words=360):
    """Expand a paragraph to approximately target_words using a large-model API.

    Returns expanded text or "error" on failure.
    """
    response = requests.post(
        url="https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {Llama_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": "meta-llama/llama-3.3-70b-instruct:free",
            "messages": [
                {
                    "role": "user",
                    "content": (
                        f"Please expand this academic paragraph to approximately "
                        f"{target_words} words:\n\n{paragraph}\n"
                        "No suggestions at the end. No word count."
                    )
                }
            ]
        },
        timeout=30
    )
    try:
        data = response.json()

        return (
            data.get("choices", [{}])[0]
                .get("message", {})
                .get("content", "⚠️ Expansion failed. Please try again.")
        )
    except:
        return "error"







# def expand_grok_ai(paragraph, target_words=360):
#     if not GROK_KEY:
#         return paragraph
#     url = "https://openrouter.ai/api/v1/chat/completions"
#     headers = {"Authorization": f"Bearer {GROK_KEY}", "Content-Type": "application/json"}
#     data = {
#         "model": "x-ai/grok-4.1-fast:free",
#         "messages": [
#             {"role": "user", "content": f"Please expand this  academic paragraph to approximately {target_words} words:\n\n{paragraph}. no suggestion at the end. no word counts please."}
#         ],
#         "extra_body": {"reasoning": {"enabled": True}}
#     }
#     try:
#         response = requests.post(url, headers=headers, json=data, timeout=15)
#         result = response.json()
#         content = result["choices"][0]["message"]["content"]
#         return content.replace("*","").replace("#","").replace("-","")
#     except Exception as e:
#         return f"Error: {e}"

# def generate_deepseek(prompt):
#     response = requests.post(
#     url="https://openrouter.ai/api/v1/chat/completions",
#     headers={
#         "Authorization": f"Bearer {DEEPSEEK_KEY}",
#         "Content-Type": "application/json",
#         "HTTP-Referer": "<YOUR_SITE_URL>", # Optional. Site URL for rankings on openrouter.ai.
#         "X-Title": "<YOUR_SITE_NAME>", # Optional. Site title for rankings on openrouter.ai.
#     },
#     data=json.dumps({
#         "model": "meta-llama/llama-guard-4-12b:free",
#         "messages": [
#         {
#             "role": "user",
#             "content": [
#             {
#                 "type": "text",
#                 "text": f"write a clear academic text{prompt}"
#             },
#             ]
#         }
#         ]
#     })
#     )

#     #resp = requests.post(url, headers=headers, json=payload, timeout=20)
#     result = response.json()

#     if "choices" not in result:
#         return f"[API Error: {result}]"

#     return result["choices"][0]["message"]["content"]



# def expand_grok_ai(paragraph, target_words=360):
#     url = "https://openrouter.ai/api/v1/chat/completions"
#     headers = {
#         "Authorization": f"Bearer {GROK_KEY}",
#         "Content-Type": "application/json",
#         "X-Title": "MyTelegramBot"
#     }

#     data = {
#         "model": "google/gemma-3-4b-it:free",
#         "messages": f"expand these paragraphs as long as {target_words}. \n {paragraph}",
#         "response_format": {"type": "text"}
#     }

#     response = requests.post(url, headers=headers, json=data, timeout=20)
#     result = response.json()

#     if "choices" not in result:
#         return f"[API Error: {result}]"

#     return result["choices"][0]["message"]["content"]

# ------------------------------
# 4. ADD CONTENT
# ------------------------------
def add_reja_page(document, plan_items, topic, lang):
    """Insert the topic and REJA page into the document using translated labels."""
    t = TRANSLATIONS[lang]
    document.add_paragraph(f"{t["mavzu_titul"]}:{topic}")
    document.add_paragraph("")
    document.add_paragraph(t["reja_titul"])
    for item in plan_items:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_page_break()
    return document


def add_content_pages(document, plan_items, paragraphs):
    """Append each plan item title and its corresponding paragraph to the document."""
    for title, para in zip(plan_items, paragraphs):
        # Add paragraph title and make it bold
        p_title = document.add_paragraph(title)
        run = p_title.runs[0] if p_title.runs else p_title.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

        # Add content paragraph
        p = document.add_paragraph(para)
        p.paragraph_format.first_line_indent = Cm(0.42)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add page break after each point
        document.add_page_break()
    return document


#this only addes the generated conclusion to the file
def add_conclusion(document, conclusion, lang):
    """Add the generated conclusion section (translated) to the document."""
    t = TRANSLATIONS[lang]
    p_xulosa = document.add_paragraph(t["xulosa_titul"])
    p_xulosa.bold = True
    p_xulosa.alignment = WD_ALIGN_PARAGRAPH.CENTER    
    for item in conclusion:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_page_break()

    return document


def add_reference(document, references, lang):
    """Append references section using translated title."""
    t = TRANSLATIONS[lang]
    document.add_paragraph(t["adabiyotlar_titul"])
    for item in references:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_page_break()
    return document

# ------------------------------
# 5. MAIN
# ------------------------------
def make_mustaqil(university, student, teacher, language, topic, pages, file_path):
    # university = input("Universitetingiz: ")
    # student = input("Bajaruvchi: ")
    # teacher = input("Qabul qiluvchi (darajasi bilan yozing masalan prof. yoki dots.'): ")
    document = make_first_page(university, student, teacher, language)
    # language = input("Qaysi tilda bo'lsin (O'zbek=uz, Ingiliz=eng, rus tili=ru): ")
    # flavour = input("Qo'shimcha talablaringizni yozing: ")

    # topic = input("Mavu: ")
    # pages = int(input("Necha sahifalik referat bo‘lsin? (1 sahifa = 300 words): "))
    
    total_words = int(pages) * 280
    
    #print("generating reja...")
    plan_prompt = f"""
    Generate a structured plan (REJA) with maximum 5 points for a referat on '{topic}'.
    - Output only 5 concise points.
    - important! user also wanted it to be in {language} language.
    - do not add the word "Kirish" and " Xulosa" or in english "preface" and "conclusion" as well as russian.
    """
    plan_text = generate_deepseek(plan_prompt)
    plan_text = plan_text.replace("*", "").replace("#", "").replace("-", "").replace("**", "").replace("##", "").replace("###", "")
    plan_items = [line.strip() for line in plan_text.split("\n") if line.strip()][:5]
    
    document = add_reja_page(document, plan_items, topic, language)

    words_per_point = total_words // len(plan_items)


    #print("generating paragraphs...")
    initial_paragraphs = []
    for point in plan_items:
        paragraph_prompt = f"Write a concise academic paragraph in {language} about: {point}. no suggestion at the end. no word counts please. write like a student."
        short_para = generate_deepseek(paragraph_prompt)
        initial_paragraphs.append(short_para)
        time.sleep(0.6)

    #print("expanding paragraphs...")
    final_paragraphs = []
    for para in initial_paragraphs:
        expanded = expand_grok_ai(para, target_words=words_per_point)
        expanded = expanded.replace("*", "").replace("#", "").replace("-", "").replace("**", "").replace("##", "").replace("###", "")
        final_paragraphs.append(expanded)
        time.sleep(0.5)

    #generate conclusion
    conclusion = []
    prompt_conclusion = f""" Generate a conclusion about these texts you wrote here {final_paragraphs}. it must be 100-150 words. texts only."""
    concluded = generate_deepseek(prompt_conclusion)
    concluded = concluded.replace("*", "").replace("#", "").replace("-", "")
    conclusion.append(concluded)

    #making ready the references
    references = []
    prompt_reference = f"""Generate a list of 7 references that you have used to describe these points {plan_text}. -output only 7 references no extra text. - important! user also wanted it to be in {language} language"""
    refer = generate_deepseek(prompt_reference)
    refer = refer.replace("*", "").replace("#", "").replace("-", "").replace("**", "").replace("##", "").replace("###", "")
    references.append(refer)  

    

    document = add_content_pages(document, plan_items, final_paragraphs)
    document = add_conclusion(document, conclusion, language)
    document = add_reference(document, references, language)

    #print('files succesfuly saved👍')
    document.save(file_path)

def run_snip(university, student, teacher, language, topic, pages, file_path):
    """Check that AI services are reachable and create the 'mustaqil' document; return "later" if services fail."""
    text1 = generate_deepseek("hello")
    text2 = expand_grok_ai("hello", 5)
    
    if text1 == "error":
        return "later"

    if text2 == 'error':
        return 'later'
    
    make_mustaqil(university, student, teacher, language, topic, pages, file_path)

def main():
    pass



if __name__ == "__main__":
    main()
