from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from dotenv import load_dotenv
import requests, json, time
from text import TRANSLATIONS
from services.models import APIKeyManager, ModelManager

DB_path = "db/models.db"
# ------------------------------
# 1. CREATE FIRST PAGE
# ------------------------------
def make_first_page(universitet, student_name, qabul_qiluvchi, lang):
    """Create and return the cover/title page as a Document object.

    Args:
        universitet: university name string
        student_name: student name string
        qabul_qiluvchi: advisor/teacher string
        lang: language code for translations
    Returns:
        A python-docx Document with the first page prepared.
    """
    t = TRANSLATIONS[lang]
    document = Document("template.docx")
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(14)

    p = document.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph(f"{universitet}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        document.add_paragraph("")

    p = document.add_paragraph()
    run = p.add_run(t["Referat_titul"])
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        document.add_paragraph("")

    bajardi = document.add_paragraph()
    run = bajardi.add_run(t["bajaruvchi_titul"])
    run.bold = True
    bajardi.add_run(f"                                                                    {student_name}")
    for _ in range(3):
        document.add_paragraph("")

    qabul_qildi = document.add_paragraph()
    run = qabul_qildi.add_run(t["qabul_qiluvchi_titul"])
    run.bold = True
    qabul_qildi.add_run(f"                                                              {qabul_qiluvchi}")
    document.add_paragraph("")
    document.add_paragraph("")

    year = time.strftime("%Y")
    yil = document.add_paragraph(f"{year}y")
    yil.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()
    return document

# ------------------------------
# 2. LOAD API KEYS
# ------------------------------
load_dotenv()
DEEPSEEK_KEY = os.getenv("REFERAT_KEY")
GROK_KEY = os.getenv("REFERAT_KEY2")

# ------------------------------
# 3. AI TEXT GENERATOR
# ------------------------------
def call_openrouter(model, api_key, prompt):
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=30
    )

    return response


# def generate_deepseek(prompt):
#     """Call Deepseek API to generate an academic paragraph for the given prompt.

#     Returns the response string or "error" on failure.
#     """
#     if not DEEPSEEK_KEY:
#         return "DEEPseek API key is not working"
#     url = "https://openrouter.ai/api/v1/chat/completions"
#     headers = {"Authorization": f"Bearer {DEEPSEEK_KEY}", "Content-Type": "application/json"}
#     data = {
#         "model": "deepseek/deepseek-chat",
#         "messages": [
#             {"role": "system", "content": "You write clear academic texts."},
#             {"role": "user", "content": f"{prompt} Language: uzbek"}
#         ]
#     }
#     try:
#         response = requests.post(url, headers=headers, data=json.dumps(data), timeout=12)
#         result = response.json()
#         return result["choices"][0]["message"]["content"]
#     except:
#         return "error"

def call_openrouter(model, api_key, prompt):
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=30
    )

    return response


def generate_deepseek(prompt):
    """
    Generate academic paragraph using DeepSeek model via OpenRouter.
    Returns generated text or error message.
    """

    model_manager = ModelManager(DB_path)
    key_manager = APIKeyManager(DB_path)

    while True:
        model = model_manager.get_next_model()
        key = key_manager.get_next_key()

        if not model or not key:
            return "No models or API keys left."

        response = call_openrouter(
            model=model,
            api_key=key,
            prompt=(
                "You write clear academic texts.\n\n"
                f"{prompt}\n\n"
                "Language: Uzbek"
            )
        )

        status = response.status_code
        print(status)

        if status == 200:
            try:
                data = response.json()
                return (
                    data.get("choices", [{}])[0]
                        .get("message", {})
                        .get("content", "⚠️ Generation failed.")
                )
            except Exception:
                return "error"

        if status in [400, 401, 429]:
            key_manager.remove_broken_key(key)
            continue

        if status in [404, 503]:
            model_manager.remove_broken_model(model)
            continue

        return "Unexpected error occurred."


def expand_grok_ai(user_text, target_words=360):
    model_manager = ModelManager(DB_path)
    key_manager = APIKeyManager(DB_path)

    prompt = (
        f"Please expand the following academic paragraph to approximately "
        f"{target_words} words.\n\n"
        f"{user_text}\n\n"
        "Do not include suggestions. Do not include word count."
    )

    while True:
        model = model_manager.get_next_model()
        key = key_manager.get_next_key()

        if not model or not key:
            return "No models or API keys left."

        response = call_openrouter(model, key, prompt)

        status = response.status_code
        print(status)

        if status == 200:
            try:
                data = response.json()
                return (
                    data.get("choices", [{}])[0]
                        .get("message", {})
                        .get("content", "⚠️ Expansion failed. Please try again.")
                )
            except Exception:
                return "error"

        if status in [400, 401, 429]:
            key_manager.remove_broken_key(key)
            continue

        if status in [404, 503]:
            model_manager.remove_broken_model(model)
            continue

        return "Unexpected error occurred."







# def expand_grok_ai(paragraph, target_words=360):
#     """Use a large-model API to expand a short paragraph to roughly target_words.

#     Args:
#         paragraph: source paragraph string
#         target_words: approximate target word count for expansion

#     Returns:
#         Expanded paragraph string, or "error" on failure.
#     """
#     response = requests.post(
#         url="https://openrouter.ai/api/v1/chat/completions",
#         headers={
#             "Authorization": f"Bearer {GROK_KEY}",
#             "Content-Type": "application/json",
#         },
#         json={
#             "model": "meta-llama/llama-3.3-70b-instruct:free",
#             "messages": [
#                 {
#                     "role": "user",
#                     "content": (
#                         f"Please expand this academic paragraph to approximately "
#                         f"{target_words} words:\n\n{paragraph}\n"
#                         "No suggestions at the end. No word count."
#                     )
#                 }
#             ]
#         },
#         timeout=30
#     )
#     try:
#         data = response.json()

#         return (
#             data.get("choices", [{}])[0]
#                 .get("message", {})
#                 .get("content", "⚠️ Expansion failed. Please try again.")
#         )
#     except:
#         return "error"
#this is the originala function below 
# def expand_grok_ai(paragraph, target_words=360):
#     if not GROK_KEY:
#         return "GROK API key is not working"
#     url = "https://openrouter.ai/api/v1/chat/completions"
#     headers = {"Authorization": f"Bearer {GROK_KEY}", "Content-Type": "application/json"}
#     data = {
#         "model": "x-ai/grok-4.1-fast:free",
#         "messages": [
#             {"role": "user", "content": f"Please expand this Uzbek academic paragraph to approximately {target_words} words:\n\n{paragraph}. no suggestion at the end. no word counts please."}
#         ],
#         "extra_body": {"reasoning": {"enabled": True}}
#     }
#     try:
#         response = requests.post(url, headers=headers, data=json.dumps(data), timeout=15)
#         result = response.json()
#         content = result["choices"][0]["message"]["content"]
#         return content.replace("*","").replace("#","").replace("-","")
#     except:
#         return "error"

# ------------------------------
# 4. ADD CONTENT
# ------------------------------
def add_reja_page(document, plan_items):
    """Add a table-of-contents (REJA) page to the document.

    Args:
        document: python-docx Document to modify
        plan_items: iterable of strings to add as plan entries

    Returns:
        The modified Document object.
    """
    document.add_paragraph("REJA:")
    # Iterate through plan points and add them formatted to the document
    for item in plan_items:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_page_break()
    return document

def add_content_pages(document, plan_items, paragraphs):
    """Add content pages: pair each plan title with its paragraph and append to the document."""
    for title, para in zip(plan_items, paragraphs):
        # Add paragraph title and make it bold
        p_title = document.add_paragraph(title)
        run = p_title.runs[0] if p_title.runs else p_title.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

        # Add content paragraph
        p = document.add_paragraph(para)
        p.paragraph_format.first_line_indent = Cm(0.42)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add page break after each point
        document.add_page_break()
    return document


#this only addes the generated conclusion to the file
def add_conclusion(document, conclusion):
    p_xulosa = document.add_paragraph("Xulosa")
    p_xulosa.bold = True
    p_xulosa.alignment = WD_ALIGN_PARAGRAPH.CENTER    
    for item in conclusion:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_page_break()

    return document


def add_reference(document, references):
    document.add_paragraph("Foydalanilgan Adabbiyotlar:")
    for item in references:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_page_break()
    return document

# ------------------------------
# 5. MAIN
# ------------------------------
def make_referat(university, student, teacher, language, topic, pages, file_path):
    """Generate a referat document by composing the title page, plan, content, conclusion and references, then save it.

    Args:
        university, student, teacher: strings for title page
        language: language code
        topic: referat topic
        pages: number of pages requested (each page ≈ 300 words)
        file_path: filepath to save the resulting .docx
    """
    document = make_first_page(university, student, teacher, language)
    total_words = int(pages) * 300

    #print("generating reja...")
    plan_prompt = f"""
    Generate a structured plan (REJA) with maximum 5 points for a referat on '{topic}'.
    - Output only 5 concise points.
    - important! user also wanted it to be in {language} language.

    """
    plan_text = generate_deepseek(plan_prompt)
    plan_text = plan_text.replace("*", "").replace("#", "").replace("-", "").replace("**", "").replace("##", "").replace("###", "")
    plan_items = [line.strip() for line in plan_text.split("\n") if line.strip()][:5]
    
    document = add_reja_page(document, plan_items)

    words_per_point = total_words // len(plan_items)


    #print("generating paragraphs...")
    initial_paragraphs = []
    for point in plan_items:
        paragraph_prompt = f"Write a concise academic paragraph in {language} about: {point}. no suggestion at the end. no word counts please. write like a student."
        short_para = generate_deepseek(paragraph_prompt)
        initial_paragraphs.append(short_para)
        time.sleep(0.6)

    #print("expanding paragraphs...")
    final_paragraphs = []
    for para in initial_paragraphs:
        expanded = expand_grok_ai(para, target_words=words_per_point)
        expanded = expanded.replace("*", "").replace("#", "").replace("-", "").replace("**", "").replace("##", "").replace("###", "")
        final_paragraphs.append(expanded)
        time.sleep(0.5)

    #generate conclusion
    conclusion = []
    prompt_conclusion = f""" Generate a conclusion about these texts you wrote here {final_paragraphs}. it must be 100-150 words. texts only."""
    concluded = generate_deepseek(prompt_conclusion)
    concluded = concluded.replace("*", "").replace("#", "").replace("-", "")
    conclusion.append(concluded)

    #making ready the references
    references = []
    prompt_reference = f"""Generate a list of 7 references that you have used to describe these points {plan_text}. -output only 7 references no extra text. - important! user also wanted it to be in {language} language"""
    refer = generate_deepseek(prompt_reference)
    refer = refer.replace("*", "").replace("#", "").replace("-", "").replace("**", "").replace("##", "").replace("###", "")
    references.append(refer)  

    

    document = add_content_pages(document, plan_items, final_paragraphs)
    document = add_conclusion(document, conclusion)
    document = add_reference(document, references)

    #print('files succesfuly saved👍')
    document.save(file_path)

def run_code(university, student, teacher, language, topic, pages, file_path):
    """Sanity-check the external AI services and then call make_referat to build the document.

    Returns "later" when remote services are unavailable so caller can retry later.
    """
    text1 = generate_deepseek("hello")
    text2 = expand_grok_ai("hello", 5)
    
    if text1 == "error":
        print("Error in deepseek")
        return "later"

    if text2 == 'error':
        print("Error in grok")
        return 'later'
    
    make_referat(university, student, teacher, language, topic, pages, file_path)

def main():
    pass

if __name__ == "__main__":
    main()
