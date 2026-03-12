from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from dotenv import load_dotenv
import requests, json, time, logging

# ------------------------------
# 1. CREATE FIRST PAGE
# ------------------------------
def make_first_page(universitet, student_name, qabul_qiluvchi):
    """Create the first/title page for a referat using the local template and return the Document."""
    document = Document("template.docx")
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(14)

    p = document.add_paragraph("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM FAN VA INNOVATSIYALAR VAZIRLIGI")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph(f"{universitet}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        document.add_paragraph("")

    p = document.add_paragraph()
    run = p.add_run('REFERAT')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        document.add_paragraph("")

    bajardi = document.add_paragraph()
    run = bajardi.add_run("Bajardi:")
    run.bold = True
    bajardi.add_run(f"                                                                    {student_name}")
    for _ in range(3):
        document.add_paragraph("")

    qabul_qildi = document.add_paragraph()
    run = qabul_qildi.add_run("Qabul qildi:")
    run.bold = True
    qabul_qildi.add_run(f"                                                              {qabul_qiluvchi}")
    document.add_paragraph("")
    document.add_paragraph("")

    yil = document.add_paragraph("2025y")
    yil.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()
    return document

# ------------------------------
# 2. LOAD API KEYS
# ------------------------------
load_dotenv()
DEEPSEEK_KEY = os.getenv("DEEPSEEK_API_KEY")
GROK_KEY = os.getenv("GROK_API_KEY")

# ------------------------------
# 3. AI TEXT GENERATOR
# ------------------------------
def generate_deepseek(prompt):
    # If API key is missing, avoid remote calls and return a safe placeholder so script doesn't hang.
    if not DEEPSEEK_KEY:
        logging.warning("DEEPSEEK API key not found. Using local mock paragraph for prompt start: %s", prompt[:80])
        # Very small, neutral Uzbek placeholder that keeps the flow going when keys are missing
        return "Bu mavzuga oid qisqacha akademik monolog. Batafsil izohlar keyingi qadamda kengaytiriladi."
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {DEEPSEEK_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "deepseek/deepseek-chat",
        "messages": [
            {"role": "system", "content": "You write clear academic texts."},
            {"role": "user", "content": f"{prompt} Language: uzbek"}
        ]
    }
    # Make requests with retries and a timeout so the script doesn't hang
    for attempt in range(1, 4):
        try:
            response = requests.post(url, headers=headers, data=json.dumps(data), timeout=10)
            result = response.json()
            if not isinstance(result, dict):
                logging.warning("deepseek: unexpected response type (not dict): %s", type(result))
                continue
            if "choices" not in result:
                logging.warning("deepseek: missing choices in response: %s", result)
                continue
            return result["choices"][0]["message"]["content"]
        except requests.exceptions.Timeout:
            logging.warning("deepseek: request timed out (attempt %d)", attempt)
        except Exception as e:
            logging.warning("deepseek: request failed (attempt %d): %s", attempt, e)
        time.sleep(attempt * 1.2)

    # fallback: give a graceful default so downstream processing continues
    logging.error("deepseek: all attempts failed for prompt starting: %s", prompt[:120])
    return "[Error generating paragraph — remote API unavailable]"

def expand_grok_ai(paragraph):
    # If the Grok key is not set, return paragraph unchanged (no remote expansion)
    if not GROK_KEY:
        logging.warning("GROK API key not found. Skipping remote expansion for paragraph start: %s", paragraph[:80])
        return paragraph
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {GROK_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "x-ai/grok-4.1-fast:free",
        "messages": [
            {"role": "user", "content": f"Please expand and elaborate this paragraph in Uzbek academically:\n\n{paragraph}"}
        ],
        "extra_body": {"reasoning": {"enabled": True}}
    }
    for attempt in range(1, 4):
        try:
            response = requests.post(url, headers=headers, data=json.dumps(data), timeout=12)
            result = response.json()
            if "choices" not in result:
                logging.warning("grok: missing choices in response (attempt %d): %s", attempt, result)
                continue
            return result["choices"][0]["message"]["content"]
        except requests.exceptions.Timeout:
            logging.warning("grok: request timed out (attempt %d)", attempt)
        except Exception as e:
            logging.warning("grok: request failed (attempt %d): %s", attempt, e)
        time.sleep(attempt * 1.0)

    logging.error("grok: all attempts failed for paragraph starting: %s", paragraph[:120])
    # return the original paragraph when expansion fails
    return paragraph

# ------------------------------
# 4. ADD REJA AND CONTENT
# ------------------------------
def add_reja_page(document, plan_items):
    """Add 'REJA' (plan) page and insert the plan items formatted into the document."""
    document.add_paragraph("REJA:")
    for item in plan_items:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_page_break()
    return document

def add_content_pages(document, paragraphs):
    """Append each paragraph as a new page to the document, applying thesis formatting."""
    for para in paragraphs:
        if para.strip():
            document.add_page_break()
            p = document.add_paragraph(para.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return document

# ------------------------------
# 5. MAIN
# ------------------------------
if __name__ == "__main__":
    # setup logging so we can see what's happening when network ops fail
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
    document = make_first_page("My University", "Azam", "Professor X")

    topic = "O'zbekiston Respublikasiga chegaradosh davlatlar"
    print("generating reja...")
    # Step 1: generate REJA + initial points
    plan_prompt = f"""
    Generate a structured plan (REJA) for a referat on '{topic}'.
    - Format: 1. point 1, 2. point 2, etc.
    - Each point should be concise.
    - Output only the list.
    """
    plan_text = generate_deepseek(plan_prompt)
    plan_items = [line.strip() for line in plan_text.split("\n") if line.strip()]
    document = add_reja_page(document, plan_items)

    print("genertating paragraphs....")
    # Step 2: generate initial paragraph for each point
    initial_paragraphs = []
    for i, point in enumerate(plan_items, 1):
        # make prompt explicit about language and academic style
        paragraph_prompt = (
            f"Write a concise academic paragraph in Uzbek (O'zbek) about the following short point: {point}."
            " Keep it formal and suitable for an academic referat."
        )
        logging.info("generating initial paragraph %d/%d for point: %s", i, len(plan_items), point[:80])
        short_para = generate_deepseek(paragraph_prompt)
        # If the generator returned an obvious error placeholder, generate a safe fallback paragraph
        if short_para.strip().startswith("[Error"):
            logging.info("using fallback paragraph for point: %s", point[:80])
            short_para = f"{point}. Bu mavzu haqida qisqacha akademik izoh: muhim jihatlar va umumiy kontekstni keltiring."  # Uzbek placeholder
        initial_paragraphs.append(short_para)
        # small pause to avoid hitting rate-limits
        time.sleep(0.6)

    print("expanding each paragraph...")
    # Step 3: expand each paragraph with Grok AI
    final_paragraphs = []
    for idx, para in enumerate(initial_paragraphs, 1):
        logging.info("expanding paragraph %d/%d…", idx, len(initial_paragraphs))
        expanded_para = expand_grok_ai(para)
        # remove extra marks like *, #, etc
        expanded_para = expanded_para.replace("*", "").replace("#", "").replace("-", "")
        final_paragraphs.append(expanded_para)
        time.sleep(0.5)
    
    print("getting ready the file...")
    # Step 4: add content to document
    document = add_content_pages(document, final_paragraphs)

    document.save("referat_final.docx")
    print("Done! Document saved as 'referat_final.docx'")
