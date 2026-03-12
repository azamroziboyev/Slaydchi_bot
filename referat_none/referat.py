from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.oxml.ns import qn
import os
from openai import OpenAI
from dotenv import load_dotenv
import requests, json




def make_first_page(universitet, student_name, qabul_qiluvchi):
    """Create and return the first/title page using `template.docx` and provided metadata."""
    document = Document("template.docx")
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # for East Asian / Word compatibility
    style.font.size = Pt(14)

    # titul beti ramkalik
    p = document.add_paragraph("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM FAN VA INNOVATSIYALAR VAZIRLIGI")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph(f"{universitet}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        document.add_paragraph("")
    

    p = document.add_paragraph()
    run = p.add_run('REFERAT')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run.bold = True
    font =run.font
    font.name = 'Times new Roman'
    font.size = Pt(56)
    for _ in range(4):
        document.add_paragraph("")


    bajardi = document.add_paragraph()
    #TODO shu yerga bajaruvchi yozilishi kerak
    run = bajardi.add_run("Bajardi:")

    run.bold = True
    bajardi.add_run(f"                                                                    {student_name}")
    for _ in range(3):
        document.add_paragraph("")

    
    qabul_qildi = document.add_paragraph()
    run = qabul_qildi.add_run("Qabul qildi:")
    run.bold = True
    qabul_qildi.add_run(f"                                                              {qabul_qiluvchi}")
    document.add_paragraph("")
    document.add_paragraph("")

    yil = document.add_paragraph("2025y")
    yil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #BU birinchi sahifa oxiri 
    document.add_page_break()
    return document

load_dotenv()

API_KEY = os.getenv("DEEPSEEK_API_KEY")  # or paste directly


# ------------------------------
# 1. AI TEXT GENERATOR (OPENROUTER)
# ------------------------------
def generate_referat(topic):
    """Call the remote model to generate referat text for the given topic and return the generated content string."""
    url = "https://openrouter.ai/api/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": "deepseek/deepseek-chat",
        "messages": [
            {"role": "system", "content": "You write clear academic texts."},
            {"role": "user", "content": f"{topic} language: uzbek"}
        ]
    }

    response = requests.post(url, headers=headers, data=json.dumps(data))
    result = response.json()

    if "choices" not in result:
        print("API ERROR:", result)
        return "Error occurred while generating."

    return result["choices"][0]["message"]["content"]


# ------------------------------
# 2. ADD SECOND PAGE TO DOCX
# ------------------------------
def add_reja_page(document, plan_items):
    """Insert a plan (REJA) page to the document with formatted entries."""
    document.add_page_break()
    document.add_paragraph("REJA:", style='Heading 1')
    for item in plan_items:
        if item.strip():
            p = document.add_paragraph(item.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return document

def add_content_pages(document, paragraphs):
    """Write each supplied paragraph on a new page and save the final file."""
    for para in paragraphs:
        if para.strip():
            document.add_page_break()  # start at top of page
            p = document.add_paragraph(para.strip())
            p.paragraph_format.first_line_indent = Cm(0.42)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save final document
    document.save("referat_final.docx")
    print("document saved")
    

# def add_second_page(document, content):
#     doc = document

#     # Add a page break before content
#     doc.add_page_break()

#     # Split generated text into paragraphs
#     paragraphs = content.split("\n")

#     for p in paragraphs:
#         if p.strip():
#             para = doc.add_paragraph(p.strip())

#             # 5-character indentation
#             para.paragraph_format.first_line_indent = Cm(0.42)
#             para.paragraph_format.line_spacing = 1.5
#             para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#             # Ensure Times New Roman formatting
#             run = para.runs[0]
#             run.font.name = "Times New Roman"
#             run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
#             run.font.size = Pt(14)

#     output_name = "updated_result.docx"
#     doc.save(output_name)
#     print(f"Saved as: {output_name}")


# ------------------------------
# MAIN
# ------------------------------
if __name__ == "__main__":
    document = make_first_page("My University", "Azam", "Professor X")

    topic = "O'zbekiston Respublikasiga chegaradosh davlatlar"

    plan_prompt = f"""
    Generate a structured plan (REJA) for a referat on '{topic}'.
    - Format as numbered list: 1. paragraph 1, 2. paragraph 2, etc.
    - Each point should be concise.
    - Maximum total words: 50-100.
    - Output only the list, no extra text.
    """
    plan_text = generate_referat(plan_prompt)
    plan_items = plan_text.split("\n")
    num_pages = 5
    total_words = num_pages * 360

    content_prompt = f"""
    Write a referat on '{topic}' in Uzbek.
    - Total words: approx {total_words}.
    - Follow the plan points from REJA.
    - Each paragraph should be one point from the plan.
    - Start each paragraph at the top of the page.
    - First-line indentation = 5 characters.
    - Scientific / academic style.
    """
    referat_text = generate_referat(content_prompt)
    paragraphs = referat_text.split("\n\n")  # split by paragraph
    # referat_text = generate_referat(topic)

    add_content_pages(document, referat_text)

    document.save("updated_result.docx")
    print("Done!")


# def get_info():
#     universitet = input("Universitetingizni kiriting: ").upper()
#     student_name = input("Ismingizni kiriting: ").title()
#     qabul_qiluvchi = input("qabul qiluvchini kiriting: ").title()
    
#     return universitet, student_name, qabul_qiluvchi


# def main():
#     university, bajaruvchi, qabul_qiluvchi = get_info()
#     make_first_page(university, bajaruvchi, qabul_qiluvchi)
    


# if __name__ == "__main__":
#     main()


