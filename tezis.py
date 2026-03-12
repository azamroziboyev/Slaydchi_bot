import requests
from generator import generate_thesis
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from services.models import APIKeyManager, ModelManager
import os


DB_path = "db/models.db"


# def expand_grok_ai(prompt):
#     response = requests.post(
#         url="https://openrouter.ai/api/v1/chat/completions",
#         headers={
#             "Authorization": f"Bearer {API_KEY}",
#             "Content-Type": "application/json",
#         },
#         json={
#             "model": "tngtech/tng-r1t-chimera:free",
#             "messages": [
#                 {
#                     "role": "user",
#                     "content": (
#                         f"{prompt}"
#                         "No suggestions at the end. No word count."
#                     )
#                 }
#             ]
#         },
#         timeout=30
#     )
#     try:
#         data = response.json()

#         return (
#             data.get("choices", [{}])[0]
#                 .get("message", {})
#                 .get("content", "⚠️ Expansion failed. Please try again.")
#         )
#     except:
#         return "error"



def call_openrouter(model, api_key, prompt):
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=30
    )

    return response

def expand_grok_ai(user_text):
    model_manager = ModelManager(DB_path)
    key_manager = APIKeyManager(DB_path)
    while True:
        model = model_manager.get_next_model()
        key = key_manager.get_next_key()

        if not model or not key:
            return "No models or API keys left."

        response = call_openrouter(model, key, user_text)

        status = response.status_code
        print(status)
        if status == 200:
            data = response.json()
            return data["choices"][0]["message"]["content"]

        if status in [400, 401, 429]:
            key_manager.remove_broken_key(key)
            continue

        if status in [404, 503]:
            model_manager.remove_broken_model(model)
            continue

        return "Unexpected error occurred."



def add_topic_writer(topic, university, student_name):
    doc = Document()
    
    # Global Style Setting
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Topic (Centered and Bold)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(topic)
    run1.bold = True  # Must be set to True
    
    # University & Name (Right-aligned and Italic)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(f"{university}\n{student_name}")
    run2.italic = True # Must be set to True
    
    return doc

def add_thesis(doc, paragraphs):
    """Create and return a Word Document formatted as a thesis.

    Args:
        paragraphs: iterable of paragraph strings to add to the document.

    Returns:
        A python-docx Document object with formatted paragraphs and page breaks.
    """
    document = doc
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(14)

    for para in paragraphs:
        # Add and format each paragraph (indent, spacing, justification), then add a page break
        p = document.add_paragraph(para)
        p.paragraph_format.first_line_indent = Cm(0.42)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add page break after each point
        document.add_page_break()
    return document


def make_thesis(language, topic, file_path, university, student_name):
    """Generate a thesis text using AI and save it to a DOCX file.

    Args:
        language: language code for generated text (e.g., 'uz').
        topic: topic string for the thesis.
        file_path: path where the generated .docx will be saved.
    """
    if language == 'uz':
        language = "uzbek latin"
    elif language == 'ru':
        language =="Russian"
    else:
        language = 'english'

    prompt_tezis = f"""
    Your task is to write a complete thesis text based on the topic which is {topic}.
    do not write additional word counts or some add and so foth at the start or nowhere in the context. just write the text asked.
    Write the thesis in strict academic {language} (ilmiy uslub), following these rules:

    📘 1. STRUCTURE OF THE THESIS
    
    The thesis must include the following mandatory sections:

    Paragraph1: Introduction to the topic, problem statement, and relevance. (do not write this)
    Paragraph2: Scientific background and literature review. Research objective (do not write this))
    Paragraph3: Methodology and research methods used.(do not write this)
    Paragraph4: Research results and conclusions.(do not write this )
    references and sources should be included at the end of the thesis under the section "FOYDALANILGAN ADABIYOTLAR".

    📘 2. WRITING REQUIREMENTS
    
    - Write in strict academic {language}.
    - Avoid personal opinions.
    - Do NOT write theory for the sake of theory. Use factual information and practical data.
    - Use long paragraphs with smooth logical transitions.
    - Do NOT use bullet points. Write full paragraphs only.
    - Avoid AI-style wording: no “as an AI language model”, no disclaimers.

    📘 3. CONTENT REQUIREMENTS
    
    For EACH chapter:

    ▪ Explain the problem clearly  
    ▪ Provide scientific background  
    ▪ Add analytical explanations  
    ▪ Add real examples (Uzbekistan, region, industry-specific)  
    ▪ Describe methods, tools, and measurement approaches (if relevant)  
    ▪ Provide comparative analysis  
    ▪ Add tables or structured text when necessary (but not bullet points)   
    ▪ Final thesis size should match the length I specify

    📘 4. REFERENCES AND SOURCES

    At the end include:

    FOYDALANILGAN ADABIYOTLAR
    - At least 5 or 7 sources
    - Uzbek laws, presidential decrees, statistics, scientific books, international sources
    - APA or GOST-style optional

    The final thesis length should be proportional to the number of pages of 3000-4000 words max.
    do not write additional word counts or some add and so foth at the end. just write the text asked.
    """

    tezislar = []
    #generated_tezis = generate_thesis(prompt_tezis)
    generated_tezis = expand_grok_ai(prompt_tezis)
    if generate_thesis == "Unexpected error occurred.":
        return "error"
    tezislar.append(generated_tezis.replace("*", "").replace("#", "").replace("-", ""))
    
    document = add_topic_writer(topic, university, student_name)
    print("topic aded")
    document = add_thesis(document, tezislar)
    print("Tezis matni qo'shildi")
    document.save(file_path)

    return file_path



    